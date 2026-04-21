#!/usr/bin/env python3
"""Generate paper/reference.docx for pandoc --reference-doc.

This is the styling backbone for the DOCX version of the QAG technical
paper.  It enforces the same typography decisions captured in
.cursor/skills/qag-paper-typesetting/SKILL.md:

  * A4 paper, 2 cm margins on all four sides.
  * Body (``Normal``) in Calibri 11 pt (falls back to Arial/DejaVu
    Sans on systems without Calibri).
  * Headings 1-3 in Calibri, bold, QGI blue (``#2E4A8F``), scaled
    16 / 14 / 12 pt.
  * Captions italic 10 pt, centered, QGI-gray (``#4B5563``).
  * Hyperlinks underlined in QGI blue.
  * Code / ``Source Code`` style in Consolas 10 pt.
  * A ``Table Grid`` style with thin 0.5 pt borders and no shading.

The rendered ``reference.docx`` is committed alongside this script so the
build is reproducible from ``QAG-Technical-Paper.md`` alone.

Usage (from repo root)::

    /tmp/qgi-paper-venv/bin/python paper/build_reference_docx.py

Requires::

    /tmp/qgi-paper-venv/bin/pip install python-docx
"""

from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


QGI_BLUE = RGBColor(0x2E, 0x4A, 0x8F)
QGI_BLUE_SOFT = RGBColor(0x4A, 0x68, 0xB0)
QGI_GRAY = RGBColor(0x4B, 0x55, 0x61)
QGI_MUTED = RGBColor(0x6B, 0x72, 0x80)
QGI_BORDER = "C7D2E4"  # muted blue-gray border, matches PDF qgirule

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


def set_run_font(run, name: str, size_pt: int, bold: bool = False,
                 italic: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), name)


def configure_style(doc: Document, style_name: str, *, font: str,
                    size_pt: int, bold: bool = False, italic: bool = False,
                    color: RGBColor | None = None,
                    space_before_pt: float | None = None,
                    space_after_pt: float | None = None,
                    keep_with_next: bool | None = None,
                    alignment: int | None = None) -> None:
    style = doc.styles[style_name]
    style.font.name = font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), font)

    pfmt = style.paragraph_format
    if space_before_pt is not None:
        pfmt.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        pfmt.space_after = Pt(space_after_pt)
    if keep_with_next is not None:
        pfmt.keep_with_next = keep_with_next
    if alignment is not None:
        pfmt.alignment = alignment


def ensure_paragraph_style(doc: Document, name: str) -> None:
    styles = doc.styles
    if name not in [s.name for s in styles]:
        styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def ensure_character_style(doc: Document, name: str) -> None:
    styles = doc.styles
    if name not in [s.name for s in styles]:
        styles.add_style(name, WD_STYLE_TYPE.CHARACTER)


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def add_table_grid_border(style) -> None:
    elem = style.element
    tbl_pr = elem.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        elem.append(tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), QGI_BORDER)
        borders.append(b)
    tbl_pr.append(borders)


def main(out_path: str) -> None:
    doc = Document()
    set_page(doc)

    configure_style(
        doc,
        "Normal",
        font=BODY_FONT,
        size_pt=11,
        space_after_pt=6,
    )
    doc.styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    doc.styles["Normal"].paragraph_format.line_spacing = 1.18

    configure_style(
        doc, "Heading 1",
        font=BODY_FONT, size_pt=18, bold=True, color=QGI_BLUE,
        space_before_pt=18, space_after_pt=6, keep_with_next=True,
    )
    configure_style(
        doc, "Heading 2",
        font=BODY_FONT, size_pt=14, bold=True, color=QGI_BLUE_SOFT,
        space_before_pt=14, space_after_pt=4, keep_with_next=True,
    )
    configure_style(
        doc, "Heading 3",
        font=BODY_FONT, size_pt=12, bold=True, color=QGI_GRAY,
        space_before_pt=10, space_after_pt=3, keep_with_next=True,
    )

    configure_style(
        doc, "Title",
        font=BODY_FONT, size_pt=26, bold=True, color=QGI_BLUE,
        space_before_pt=0, space_after_pt=6,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    if "Subtitle" in [s.name for s in doc.styles]:
        configure_style(
            doc, "Subtitle",
            font=BODY_FONT, size_pt=14, italic=True, color=QGI_MUTED,
            space_after_pt=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    if "Author" in [s.name for s in doc.styles]:
        configure_style(
            doc, "Author",
            font=BODY_FONT, size_pt=12, color=QGI_GRAY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=3,
        )
    if "Date" in [s.name for s in doc.styles]:
        configure_style(
            doc, "Date",
            font=BODY_FONT, size_pt=11, color=QGI_MUTED,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=18,
        )

    if "Caption" in [s.name for s in doc.styles]:
        configure_style(
            doc, "Caption",
            font=BODY_FONT, size_pt=10, italic=True, color=QGI_GRAY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before_pt=3, space_after_pt=10,
        )

    ensure_paragraph_style(doc, "Source Code")
    configure_style(
        doc, "Source Code",
        font=MONO_FONT, size_pt=10, color=QGI_GRAY,
        space_before_pt=3, space_after_pt=3,
    )

    for toc_name, size in (("TOC Heading", 14), ("TOC 1", 11), ("TOC 2", 11), ("TOC 3", 11)):
        if toc_name in [s.name for s in doc.styles]:
            configure_style(
                doc, toc_name,
                font=BODY_FONT, size_pt=size,
                bold=toc_name in ("TOC Heading", "TOC 1"),
                color=QGI_BLUE if toc_name == "TOC Heading" else QGI_GRAY,
                space_before_pt=4 if toc_name != "TOC Heading" else 12,
                space_after_pt=2 if toc_name != "TOC Heading" else 6,
            )

    for pandoc_style, kwargs in (
        ("Body Text", dict(font=BODY_FONT, size_pt=11,
                           space_before_pt=0, space_after_pt=6)),
        ("First Paragraph", dict(font=BODY_FONT, size_pt=11,
                                 space_before_pt=0, space_after_pt=6)),
        ("Compact", dict(font=BODY_FONT, size_pt=11,
                         space_before_pt=0, space_after_pt=2)),
        ("Image Caption", dict(font=BODY_FONT, size_pt=10, italic=True,
                               color=QGI_GRAY, space_before_pt=3,
                               space_after_pt=10,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER)),
        ("Captioned Figure", dict(font=BODY_FONT, size_pt=11,
                                   space_before_pt=8, space_after_pt=0,
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER)),
        ("Abstract", dict(font=BODY_FONT, size_pt=11, italic=True,
                          color=QGI_GRAY,
                          space_before_pt=0, space_after_pt=6)),
        ("Abstract Title", dict(font=BODY_FONT, size_pt=12, bold=True,
                                color=QGI_BLUE_SOFT,
                                space_before_pt=18, space_after_pt=4,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)),
    ):
        ensure_paragraph_style(doc, pandoc_style)
        configure_style(doc, pandoc_style, **kwargs)

    ensure_character_style(doc, "Hyperlink")
    hyperlink_style = doc.styles["Hyperlink"]
    hyperlink_style.font.name = BODY_FONT
    hyperlink_style.font.underline = True
    hyperlink_style.font.color.rgb = QGI_BLUE

    if "Table Grid" in [s.name for s in doc.styles]:
        tg = doc.styles["Table Grid"]
        tg.font.name = BODY_FONT
        tg.font.size = Pt(10.5)
        add_table_grid_border(tg)

    out = pathlib.Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"wrote {out} ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    import sys

    default_target = pathlib.Path(__file__).parent / "reference.docx"
    target = sys.argv[1] if len(sys.argv) > 1 else str(default_target)
    main(target)
